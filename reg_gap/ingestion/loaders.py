"""
Document loaders for PDF, HTML, and DOCX regulatory texts.
"""

import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional
import re


@dataclass
class RegulatoryDocument:
    """Represents a loaded regulatory document."""
    
    content: str
    source_path: str
    jurisdiction: str
    document_type: str
    version: str = "1.0"
    effective_date: Optional[datetime] = None
    metadata: dict = field(default_factory=dict)
    
    def __post_init__(self):
        if not self.content:
            raise ValueError("Document content cannot be empty")


class DocumentLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """Load a document from the given path."""
        pass
    
    @abstractmethod
    def supports(self, path: str) -> bool:
        """Check if this loader supports the given file type."""
        pass


class PDFLoader(DocumentLoader):
    """Loader for PDF regulatory documents."""
    
    def supports(self, path: str) -> bool:
        """Check if this loader supports the given file path.

        Args:
            path: File path to check.

        Returns:
            True if the file is a PDF, False otherwise.
        """
        return path.lower().endswith('.pdf')
    
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """
        Load a PDF document.
        
        Args:
            path: Path to the PDF file
            jurisdiction: Regulatory jurisdiction (e.g., "US-SEC", "EU-MiFID")
            **kwargs: Additional metadata
            
        Returns:
            RegulatoryDocument with extracted text
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"PDF file not found: {path}")
        
        try:
            import pypdf
            
            text_parts = []
            with open(path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            content = "\n\n".join(text_parts)
            
        except ImportError:
            # Fallback: return placeholder for testing without pypdf
            content = self._fallback_load(path)
        
        return RegulatoryDocument(
            content=content,
            source_path=path,
            jurisdiction=jurisdiction,
            document_type="pdf",
            version=kwargs.get("version", "1.0"),
            effective_date=kwargs.get("effective_date"),
            metadata=kwargs.get("metadata", {})
        )
    
    def _fallback_load(self, path: str) -> str:
        """Fallback when pypdf is not available.

        Args:
            path: Path to the PDF file.

        Returns:
            Placeholder string indicating the PDF source path.
        """
        return f"[PDF content from: {path}]"


class HTMLLoader(DocumentLoader):
    """Loader for HTML regulatory documents."""
    
    def supports(self, path: str) -> bool:
        """Check if this loader supports the given file path.

        Args:
            path: File path to check.

        Returns:
            True if the file is an HTML file, False otherwise.
        """
        return path.lower().endswith(('.html', '.htm'))
    
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """
        Load an HTML document.
        
        Args:
            path: Path to the HTML file
            jurisdiction: Regulatory jurisdiction
            **kwargs: Additional metadata
            
        Returns:
            RegulatoryDocument with extracted text
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"HTML file not found: {path}")
        
        with open(path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for element in soup(['script', 'style', 'nav', 'footer', 'header']):
                element.decompose()
            
            content = soup.get_text(separator='\n', strip=True)
            
        except ImportError:
            # Fallback: basic HTML tag stripping
            content = self._strip_html_tags(html_content)
        
        return RegulatoryDocument(
            content=content,
            source_path=path,
            jurisdiction=jurisdiction,
            document_type="html",
            version=kwargs.get("version", "1.0"),
            effective_date=kwargs.get("effective_date"),
            metadata=kwargs.get("metadata", {})
        )
    
    def _strip_html_tags(self, html: str) -> str:
        """Basic HTML tag stripping without BeautifulSoup.

        Args:
            html: Raw HTML content string.

        Returns:
            Plain text with HTML tags, scripts, and styles removed.
        """
        clean = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=re.DOTALL | re.IGNORECASE)
        clean = re.sub(r'<[^>]+>', ' ', clean)
        clean = re.sub(r'\s+', ' ', clean)
        return clean.strip()


class DOCXLoader(DocumentLoader):
    """Loader for DOCX regulatory documents."""
    
    def supports(self, path: str) -> bool:
        """Check if this loader supports the given file path.

        Args:
            path: File path to check.

        Returns:
            True if the file is a DOCX file, False otherwise.
        """
        return path.lower().endswith('.docx')
    
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """
        Load a DOCX document.
        
        Args:
            path: Path to the DOCX file
            jurisdiction: Regulatory jurisdiction
            **kwargs: Additional metadata
            
        Returns:
            RegulatoryDocument with extracted text
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"DOCX file not found: {path}")
        
        try:
            from docx import Document
            
            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            
        except ImportError:
            content = self._fallback_load(path)
        
        return RegulatoryDocument(
            content=content,
            source_path=path,
            jurisdiction=jurisdiction,
            document_type="docx",
            version=kwargs.get("version", "1.0"),
            effective_date=kwargs.get("effective_date"),
            metadata=kwargs.get("metadata", {})
        )
    
    def _fallback_load(self, path: str) -> str:
        """Fallback when python-docx is not available.

        Args:
            path: Path to the DOCX file.

        Returns:
            Placeholder string indicating the DOCX source path.
        """
        return f"[DOCX content from: {path}]"


class TextLoader(DocumentLoader):
    """Loader for plain text regulatory documents."""
    
    def supports(self, path: str) -> bool:
        """Check if this loader supports the given file path.

        Args:
            path: File path to check.

        Returns:
            True if the file is a plain text file, False otherwise.
        """
        return path.lower().endswith('.txt')
    
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """Load a plain text document.

        Args:
            path: Path to the text file.
            jurisdiction: Regulatory jurisdiction (e.g., "US-SEC", "EU-MiFID").
            **kwargs: Additional metadata including version, effective_date, and metadata.

        Returns:
            RegulatoryDocument with the text content.

        Raises:
            FileNotFoundError: If the text file does not exist.
        """
        if not os.path.exists(path):
            raise FileNotFoundError(f"Text file not found: {path}")
        
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return RegulatoryDocument(
            content=content,
            source_path=path,
            jurisdiction=jurisdiction,
            document_type="txt",
            version=kwargs.get("version", "1.0"),
            effective_date=kwargs.get("effective_date"),
            metadata=kwargs.get("metadata", {})
        )


class UniversalLoader:
    """Universal loader that selects appropriate loader based on file type."""
    
    def __init__(self):
        """Initialize the UniversalLoader with all available document loaders."""
        self.loaders = [
            PDFLoader(),
            HTMLLoader(),
            DOCXLoader(),
            TextLoader(),
        ]
    
    def load(self, path: str, jurisdiction: str, **kwargs) -> RegulatoryDocument:
        """Load a document using the appropriate loader.

        Args:
            path: Path to the document file.
            jurisdiction: Regulatory jurisdiction (e.g., "US-SEC", "EU-MiFID").
            **kwargs: Additional metadata passed to the specific loader.

        Returns:
            RegulatoryDocument with extracted content.

        Raises:
            ValueError: If no loader supports the given file type.
        """
        for loader in self.loaders:
            if loader.supports(path):
                return loader.load(path, jurisdiction, **kwargs)
        
        raise ValueError(f"Unsupported file type: {path}")
    
    def load_directory(self, directory: str, jurisdiction: str, **kwargs) -> list[RegulatoryDocument]:
        """Load all supported documents from a directory.

        Args:
            directory: Path to the directory to scan.
            jurisdiction: Regulatory jurisdiction to assign to all documents.
            **kwargs: Additional metadata passed to each loader.

        Returns:
            List of RegulatoryDocument objects for all successfully loaded files.
        """
        documents = []
        path = Path(directory)
        
        for file_path in path.rglob("*"):
            if file_path.is_file():
                for loader in self.loaders:
                    if loader.supports(str(file_path)):
                        try:
                            doc = loader.load(str(file_path), jurisdiction, **kwargs)
                            documents.append(doc)
                        except Exception as e:
                            print(f"Warning: Failed to load {file_path}: {e}")
                        break
        
        return documents
